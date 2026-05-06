#!/usr/bin/env python3
"""
generate_reports.py — VAL.PEA
Génère automatiquement 2 documents chaque semaine :
  1. VENDREDI soir : Revue portefeuille (10 positions, fiches individuelles + avis)
  2. DIMANCHE matin : Signaux (Grade A + zone + score ≥ 70)
Appelé par GitHub Actions via update_screener.yml
"""

import re, json, sys, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── HELPERS DOCX ─────────────────────────────────────────────────────────────
def rgb(h): return RGBColor.from_string(h)

def setbg(cell, h):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), h)
    tcPr.append(shd)

def set_left_border(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz in [('top',4),('bottom',4),('left',20),('right',4)]:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz)); b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def rn(p, text, bold=False, color="333333", size=9, italic=False):
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = 'Arial'; r.font.size = Pt(size); r.font.color.rgb = rgb(color)
    return r

def para(doc, text, bold=False, color="333333", size=9.5, italic=False, before=2, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    rn(p, text, bold=bold, color=color, size=size, italic=italic)

def hdr(doc, text, color="0F2540", size=11, before=8, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    rn(p, text, bold=True, color=color, size=size)

def box(doc, text, bg_c="EFF6FF", col="1E3A5F", size=9):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.cell(0,0); setbg(c, bg_c)
    c.paragraphs[0].paragraph_format.space_before = Pt(5)
    c.paragraphs[0].paragraph_format.space_after = Pt(5)
    rn(c.paragraphs[0], text, size=size, color=col)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def avis_box(doc, verdict, texte, bg_c, col):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.cell(0,0); setbg(cell, bg_c); set_left_border(cell, col)
    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    rn(cell.paragraphs[0], verdict + "  ", bold=True, color=col, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(6)
    rn(p2, texte, size=9, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def two_col(doc, rows_data, h1="", h2=""):
    t = doc.add_table(rows=len(rows_data)+1, cols=2); t.style = 'Table Grid'
    for j,(h,hbg) in enumerate([(h1,"0F2540"),(h2,"0F2540")]):
        c = t.cell(0,j); setbg(c, hbg)
        rn(c.paragraphs[0], h, bold=True, color="FFFFFF", size=8.5)
    for i,(c1,c2,bg1,bg2,b1,col1) in enumerate(rows_data):
        ca = t.cell(i+1,0); setbg(ca, bg1); rn(ca.paragraphs[0], c1, bold=b1, color=col1, size=9)
        cb = t.cell(i+1,1); setbg(cb, bg2); rn(cb.paragraphs[0], c2, size=9)
    return t

def four_col(doc, rows_data, headers):
    t = doc.add_table(rows=len(rows_data)+1, cols=4); t.style = 'Table Grid'
    for j,h in enumerate(headers):
        c = t.cell(0,j); setbg(c,"0F2540"); rn(c.paragraphs[0],h,bold=True,color="FFFFFF",size=8)
    for i,row in enumerate(rows_data):
        for j,(text,bg_c,bold_,col_) in enumerate(row):
            c = t.cell(i+1,j); setbg(c,bg_c); rn(c.paragraphs[0],text,bold=bold_,color=col_,size=8.5)
    return t

def sep(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top'); top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'12'); top.set(qn('w:color'),'F0D080')
    pBdr.append(top); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(8)

def garde(doc, titre, sous_titre, detail, date_str):
    for txt, sz, col, bold, ital in [
        ("VAL.PEA — " + titre, 11, "6B7B8D", True, False),
        (sous_titre, 24, "0F2540", True, False),
        (detail, 11, "1A3A5C", True, False),
        (date_str, 9, "888888", False, True),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4 if sz < 14 else 22 if sz > 20 else 8)
        r = p.add_run(txt); r.bold = bold; r.italic = ital
        r.font.name = 'Georgia' if sz > 20 else 'Arial'
        r.font.size = Pt(sz); r.font.color.rgb = rgb(col)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── PARSER HTML ──────────────────────────────────────────────────────────────
def parse_stocks(html_content):
    s_start = html_content.find("const S=[")
    s_end = html_content.find("\n];\n", s_start)
    if s_end < 0:
        s_end = html_content.find("\n];", s_start)
    stocks = []
    for m in re.finditer(r"\{ticker:'([^']+)'.*?(?=\n\n\{ticker:|\n\n\];)", 
                          html_content[s_start:s_end], re.DOTALL):
        block = m.group(0)
        ticker = m.group(1)
        def gn(k): mx=re.search(r'\b'+k+r':([-\d.]+)',block); return float(mx.group(1)) if mx else 0
        def gs(k): mx=re.search(r"\b"+k+r":'([^']*)'",block); return mx.group(1) if mx else ''
        thesis_m = re.search(r'thesis:"([^"]+)"', block)
        contra_m = re.search(r'contra:"([^"]+)"', block)
        track_raw = re.findall(r"\{y:'([^']+)',p:'([^']+)',ok:'([^']+)'\}", block)
        price = gn('price')
        el = gn('el'); eh = gn('eh'); stop = gn('stop')
        o1 = gn('o1'); dcfm = gn('dcfm')
        in_zone = (el <= price <= eh) if el and eh and price else False
        dcf_zone = (dcfm > 0 and price <= dcfm * 0.88) if dcfm else False
        upside = round((dcfm - price) / price * 100, 1) if dcfm and price else 0
        rr = round((o1 - price) / (price - stop), 1) if o1 and stop and price > stop else 0
        mm200 = gn('mm200')
        rsi = gn('rsi')
        # Score QARP simplifié pour le filtrage
        score_letter = gs('score')
        stocks.append({
            'ticker': ticker, 'name': gs('name'), 'sector': gs('sector'),
            'score': score_letter, 'price': price,
            'el': el, 'eh': eh, 'stop': stop, 'o1': o1, 'o2': gn('o2'),
            'dcfm': dcfm, 'dcfb': gn('dcfb'), 'dcfu': gn('dcfu'),
            'roe': gn('roe'), 'margin': gn('margin'), 'fcf': gn('fcf'),
            'debt': gn('debt'), 'pio': gn('pio'), 'rsi': rsi, 'mm200': mm200,
            'pe': gn('pe'), 'pb': gn('pb'), 'yield': gn('yield'),
            'revg': gn('revg'), 'marg_trend': gs('marg_trend'),
            'cb': gn('cb'), 'ch': gn('ch'), 'cs': gn('cs'),
            'in_zone': in_zone, 'dcf_zone': dcf_zone,
            'upside': upside, 'rr': rr,
            'above_mm200': price > mm200 if mm200 else False,
            'rsi_ok': 25 <= rsi <= 60 if rsi else False,
            'thesis': thesis_m.group(1) if thesis_m else '',
            'contra': contra_m.group(1) if contra_m else '',
            'track': track_raw,
        })
    return stocks

def parse_ptf(html_content):
    ptf_pos = html_content.find("PTF_VAL_DEFAULT")
    ptf_end = html_content.find("];", ptf_pos) + 2
    ptf_block = html_content[ptf_pos:ptf_end]
    positions = []
    for m in re.finditer(r"ticker:'([^']+)',qty:(\d+),pru:([\d.]+)", ptf_block):
        positions.append({
            'ticker': m.group(1),
            'qty': int(m.group(2)),
            'pru': float(m.group(3)),
        })
    return positions

# ─── CALCUL QARP SIMPLIFIÉ ────────────────────────────────────────────────────
def calc_qarp(s):
    q = r = b = v = mo = 0
    # Qualité
    if s['roe'] >= 20: q = 20
    elif s['roe'] >= 15: q = 15
    elif s['roe'] >= 10: q = 10
    else: q = 5
    # Rentabilité
    if s['margin'] >= 15 and s['fcf'] >= 8: r = 20
    elif s['margin'] >= 10 and s['fcf'] >= 5: r = 15
    elif s['margin'] >= 5: r = 10
    else: r = 5
    # Bilan
    pio_score = min(int(s['pio']) * 2, 10) if s['pio'] else 5
    debt_score = 10 if s['debt'] < 1.5 else 7 if s['debt'] < 3 else 3
    b = pio_score + debt_score
    # Valorisation
    if s['upside'] >= 30: v = 20
    elif s['upside'] >= 15: v = 15
    elif s['upside'] >= 5: v = 10
    elif s['upside'] >= 0: v = 7
    else: v = 2
    # Momentum
    zone_pt = 10 if (s['in_zone'] or s['dcf_zone']) else 0
    mm200_pt = 5 if s['above_mm200'] else 0
    rsi_pt = 5 if s['rsi_ok'] else 0
    mo = zone_pt + mm200_pt + rsi_pt
    total = q + r + b + v + mo
    return {'q':q,'r':r,'b':b,'v':v,'m':mo,'total':total}

# ─── AVIS AUTOMATIQUES ────────────────────────────────────────────────────────
def gen_avis(pos, s, qarp):
    """Génère l'avis automatique selon les règles du screener"""
    p = pos['cours']; pru = pos['pru']
    pct = pos['pct']; rr = s['rr']
    in_zone = s['in_zone'] or s['dcf_zone']
    score = s['score']

    # Stop proche
    if p <= s['stop'] * 1.05 and s['stop'] > 0:
        return ('⚠️ STOP PROCHE — distance critique',
                f"Stop à {s['stop']}€. Distance : {round((p-s['stop'])/p*100,1)}% seulement. "
                f"Vente mécanique si atteint. Ne pas attendre le lundi.",
                'FEE2E2', 'DC2626')

    # Hors zone + forte PV + upside négatif
    if not in_zone and pct > 40 and s['upside'] < 0:
        pv_half = round((p - pru) * pos['qty'] / 2)
        return ('ℹ️ PV ÉLEVÉE — upside DCF négatif',
                f"PV de {pct:.1f}% hors zone avec upside DCF négatif ({s['upside']}%). "
                f"Vendre 50% ({pos['qty']//2} titre(s)) = encaisser ~{pv_half}€. "
                f"Garder le reste pour exposition + dividende.",
                'FEE2E2', 'E05000')

    # En zone + R/R excellent + score B/A
    if in_zone and rr >= 2.0 and score in ['A','B'] and qarp['total'] >= 60:
        return ('✅ ZONE ACTIVE — R/R favorable',
                f"Cours en zone, R/R {rr}x ≥ 2x. Score {qarp['total']}/100. "
                f"Conditions réunies pour un renforcement ciblé (max +3% du PEA). "
                f"Stop à {s['stop']}€ à définir avant passage d'ordre.",
                'DCFCE7', '16A34A')

    # En zone + R/R correct
    if in_zone and 1.0 <= rr < 2.0:
        return ('🟡 ZONE ACTIVE — R/R < 2x',
                f"Cours en zone mais R/R {rr}x < 2x. "
                f"Garder la position, attendre retour vers {s['el']}€ pour R/R ≥ 2x. "
                f"Ne pas renforcer ici.",
                'FFFBEB', 'D97706')

    # Hors zone mais position saine
    if not in_zone and pct > 0 and s['upside'] > 0:
        return ('⚪ HORS ZONE',
                f"Position en PV ({pct:.1f}%). Cours hors zone. "
                f"Conserver, ne pas renforcer. Surveiller retour vers {s['el']}€.",
                'FFFBEB', 'D97706')

    # En MV sans zone
    if pct < -10 and not in_zone:
        return ('⚪ HORS ZONE — cours sous zone',
                f"Moins-value de {pct:.1f}% et cours hors zone. "
                f"Vérifier si la thèse est cassée. Si fondamentaux intacts : tenir. "
                f"Si dégradation des résultats : réduire.",
                'FEF3C7', 'D97706')

    # Par défaut
    return ('⚪ SURVEILLER',
            f"Position {'en PV' if pct >= 0 else 'en MV'} de {pct:.1f}%. "
            f"Surveiller chaque vendredi. Agir si zone atteinte avec R/R ≥ 2x.",
            'F8FAFC', '6B7280')


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 : REVUE VENDREDI SOIR
# ═══════════════════════════════════════════════════════════════════════════════
def generate_revue_vendredi(html_content, output_path, date_str):
    stocks_all = parse_stocks(html_content)
    ptf_raw = parse_ptf(html_content)

    # Enrichir les positions avec les données S[]
    portfolio = []
    for pos in ptf_raw:
        if pos['ticker'] in ('DCAM','PAEEM'):
            continue  # ETF, pas de fiche
        s = next((x for x in stocks_all if x['ticker'] == pos['ticker']), None)
        if not s:
            continue
        cours = s['price']
        pnl = round((cours - pos['pru']) * pos['qty'], 2)
        pct = round((cours - pos['pru']) / pos['pru'] * 100, 2)
        qarp = calc_qarp(s)
        portfolio.append({
            **pos,
            's': s, 'cours': cours,
            'val': round(cours * pos['qty'], 2),
            'pnl': pnl, 'pct': pct,
            'qarp': qarp,
            'rr': s['rr'],
            'above_mm200': s['above_mm200'],
            'rsi_ok': s['rsi_ok'],
        })

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    garde(doc, "REVUE DE PORTEFEUILLE",
          f"Vendredi {date_str}",
          f"{len(portfolio)} positions · Fiches individuelles · Avis détaillés",
          "Cours de clôture · Analyse géopolitique · Non-conseil")

    # Synthèse
    total_val = sum(p['val'] for p in portfolio)
    total_invest = sum(p['pru'] * p['qty'] for p in portfolio)
    total_pnl = total_val - total_invest
    total_pct = (total_val / total_invest - 1) * 100 if total_invest else 0

    para(doc,
         f"Valorisation : {total_val:.0f}€  |  Investi : {total_invest:.0f}€  |  "
         f"PV/MV : {'+' if total_pnl>=0 else ''}{total_pnl:.0f}€ ({'+' if total_pct>=0 else ''}{total_pct:.1f}%)",
         bold=True, color="0F2540", size=10, before=2, after=4)

    # Tableau synthèse
    headers = ["Ticker","Société","Qté","PRU","Cours","PV/MV %","PV/MV €","Score","Avis"]
    t = doc.add_table(rows=len(portfolio)+1, cols=9); t.style = 'Table Grid'
    for j,h in enumerate(headers):
        c = t.cell(0,j); setbg(c,"0F2540"); rn(c.paragraphs[0],h,bold=True,color="FFFFFF",size=8)
    for i, pos in enumerate(portfolio):
        avis_t, _, av_bg, av_col = gen_avis(pos, pos['s'], pos['qarp'])
        avis_short = avis_t.split('—')[0].strip().replace('✅','').replace('🟡','').replace('🔴','').replace('⚠️','').replace('🟠','').strip()
        for j,(text,bg_c,bold_,col_) in enumerate([
            (pos['ticker'],"EFF6FF",True,"0F2540"),
            (pos['s']['name'][:12],"FAFAFA",False,"333333"),
            (str(pos['qty']),"FAFAFA",False,"333333"),
            (f"{pos['pru']:.2f}","FAFAFA",False,"333333"),
            (f"{pos['cours']:.2f}","FAFAFA",False,"333333"),
            (f"{'+' if pos['pct']>=0 else ''}{pos['pct']:.1f}%","DCFCE7" if pos['pct']>=0 else "FEE2E2",True,"16A34A" if pos['pct']>=0 else "DC2626"),
            (f"{'+' if pos['pnl']>=0 else ''}{pos['pnl']:.0f}€","DCFCE7" if pos['pnl']>=0 else "FEE2E2",True,"16A34A" if pos['pnl']>=0 else "DC2626"),
            (f"{pos['s']['score']} {pos['qarp']['total']}/100","F8FAFC",True,"0F2540"),
            (avis_short[:12], av_bg, True, av_col),
        ]):
            c = t.cell(i+1,j); setbg(c,bg_c); rn(c.paragraphs[0],text,bold=bold_,color=col_,size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


    # ── SECTION NICOLAS CHÉRON ─────────────────────────────────
    hdr(doc, "Analyse de marché — Nicolas Chéron (bimensuel)", "1D4ED8", size=11, before=6, after=3)
    box(doc,
        "NOTE : Chéron publie son décryptage tous les 15 jours sur YouTube (Centrale des Marchés). "
        "Il analyse les indices (CAC40, S&P500, DAX), les secteurs et 3-5 actions spécifiques avec niveaux techniques précis. "
        "Intégrer ses niveaux de résistance/support dans la surveillance des positions.",
        "EFF6FF", "1D4ED8", 9)
    
    # Tableau des points clés Chéron à surveiller cette semaine
    two_col(doc, [
        ("CAC40", "Surveiller support 7800 pts · Résistance 8200 pts", "F8FAFC","F8FAFC",True,"0F2540"),
        ("S&P500", "Tendance haussière long terme intacte · Support 5000 pts", "F8FAFC","F8FAFC",True,"0F2540"),
        ("EUR/USD", "Zone 1.05-1.10 · Impact sur les positions exportatrices (AIR)", "F8FAFC","F8FAFC",True,"0F2540"),
        ("Taux 10 ans FR", "< 3.5% = favorable pour la valorisation des actions de croissance", "F8FAFC","F8FAFC",True,"0F2540"),
        ("Actions à surveiller", "Consulter la dernière vidéo pour les niveaux spécifiques", "EFF6FF","EFF6FF",True,"1D4ED8"),
    ], "Indicateur (Chéron)", "Analyse bimensuelle")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── ALLOCATION POST-VENTE TTE ──────────────────────────────
    # Calculer si des ventes ont généré du cash à réallouer
    cash_from_sales = sum(
        (p['cours'] - p['pru']) * p['qty'] 
        for p in portfolio 
        if p['pct'] > 40 and not (p['s']['in_zone'] or p['s']['dcf_zone'])
    )
    if cash_from_sales > 200:
        # Trouver les meilleures opportunités de réallocation
        best_opps = [p for p in portfolio if (p['s']['in_zone'] or p['s']['dcf_zone']) and p['s']['rr'] >= 1.5]
        if best_opps:
            hdr(doc, "Suggestion de réallocation", "16A34A", size=10, before=4, after=2)
            opp_rows = []
            for opp in best_opps[:3]:
                s = opp['s']
                opp_rows.append((
                    opp['ticker'], 
                    f"Zone {s['el']}-{s['eh']}€ · R/R {s['rr']}x · Upside +{s['upside']}%",
                    "DCFCE7","DCFCE7",True,"16A34A"
                ))
            two_col(doc, opp_rows, f"Cash disponible ~{cash_from_sales:.0f}€", "Opportunité de réallocation")
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── FICHES INDIVIDUELLES ──
    for idx, pos in enumerate(portfolio):
        s = pos['s']
        if idx > 0:
            sep(doc)

        # Header
        th = doc.add_table(rows=1, cols=2); th.style = 'Table Grid'
        gc_map = {'A':'F0D080','B':'F0D080','C':'DDDDDD','D':'FF8888'}
        gc_c = gc_map.get(s['score'], 'DDDDDD')
        cl = th.cell(0,0); setbg(cl,"0F2540")
        p = cl.paragraphs[0]
        rn(p, f"[{s['score']}] ", bold=True, color=gc_c, size=12)
        rn(p, pos['ticker'] + "  ", bold=True, color="F0D080", size=17)
        rn(p, s['name'][:24], color="AABBCC", size=9.5)
        p2 = cl.add_paragraph()
        rn(p2, f"{s['sector']}  ·  {pos['qty']} titre(s)  ·  PRU {pos['pru']:.2f}€  ·  "
               f"Score QARP {pos['qarp']['total']}/100", color="AABBCC", size=8)
        cr = th.cell(0,1)
        setbg(cr, "1A3A5C" if pos['pnl']>=0 else "3D1515")
        p3 = cr.paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn(p3, f"{'+' if pos['pct']>=0 else ''}{pos['pct']:.1f}%  ", bold=True,
            color="F0D080" if pos['pnl']>=0 else "FF8888", size=18)
        p4 = cr.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn(p4, f"{'+' if pos['pnl']>=0 else ''}{pos['pnl']:.0f}€  |  val. {pos['val']:.0f}€",
            color="AABBCC", size=9)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Zones
        in_zone = s['in_zone'] or s['dcf_zone']
        two_col(doc, [
            ("Cours actuel", f"{pos['cours']}€", "F8FAFC","F8FAFC",True,"0F2540"),
            ("Zone achat [el–eh]",
             f"{s['el']}€ – {s['eh']}€  {'✅ cours en zone' if in_zone else '— hors zone'}",
             "DCFCE7" if in_zone else "F8FAFC","DCFCE7" if in_zone else "F8FAFC",True,
             "16A34A" if in_zone else "888888"),
            ("Stop Loss",
             f"{s['stop']}€  ({round((pos['cours']-s['stop'])/pos['cours']*100,1) if pos['cours'] else 0}% de marge)",
             "FEE2E2","FEE2E2",True,"DC2626"),
            ("Objectif 1",
             f"{s['o1']}€  (+{round((s['o1']-pos['cours'])/pos['cours']*100,1) if pos['cours'] else 0}%)",
             "DCFCE7","DCFCE7",True,"16A34A"),
            ("DCF Médian",
             f"{s['dcfm']}€  (upside {s['upside']}%)",
             "EFF6FF","EFF6FF",True,"1D4ED8"),
            ("R/R",
             f"{s['rr']}x  {'✅ favorable' if s['rr']>=1.5 else '⚠️ insuffisant'}",
             "DCFCE7" if s['rr']>=1.5 else "FFFBEB","DCFCE7" if s['rr']>=1.5 else "FFFBEB",
             True,"16A34A" if s['rr']>=1.5 else "D97706"),
        ], "Indicateur", "Valeur")
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Fondamentaux
        four_col(doc, [
            [("ROE","EFF6FF",True,"0F2540"),(f"{s['roe']}%","FAFAFA",False,"333333"),
             ("Marge nette","EFF6FF",True,"0F2540"),(f"{s['margin']}%","FAFAFA",False,"333333")],
            [("FCF yield","EFF6FF",True,"0F2540"),(f"{s['fcf']}%","FAFAFA",False,"333333"),
             ("Dette/EBITDA","EFF6FF",True,"0F2540"),(f"{s['debt']}x","FAFAFA",False,"333333")],
            [("Piotroski","EFF6FF",True,"0F2540"),(f"{int(s['pio'])}/9","FAFAFA",True,"16A34A" if s['pio']>=7 else "D97706"),
             ("PE","EFF6FF",True,"0F2540"),(f"{s['pe']}x","FAFAFA",False,"333333")],
            [("MM200","EFF6FF",True,"0F2540"),("✅ au-dessus" if s['above_mm200'] else "⚠️ en dessous","DCFCE7" if s['above_mm200'] else "FEE2E2",False,"16A34A" if s['above_mm200'] else "DC2626"),
             ("RSI","EFF6FF",True,"0F2540"),(f"{s['rsi']} {'✅' if s['rsi_ok'] else '⚠️'}","DCFCE7" if s['rsi_ok'] else "FFFBEB",False,"16A34A" if s['rsi_ok'] else "D97706")],
        ], ["Ratio","Valeur","Ratio","Valeur"])
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

        # Triptyque
        tri = [
            (f"RSI {s['rsi']}", "✅ Zone valeur (25-60)" if s['rsi_ok'] else "⚠️ Hors zone",
             "DCFCE7" if s['rsi_ok'] else "FFFBEB","DCFCE7" if s['rsi_ok'] else "FFFBEB",True,"16A34A" if s['rsi_ok'] else "D97706"),
            ("MM200", f"✅ Cours > MM200 ({s['mm200']}€)" if s['above_mm200'] else f"⚠️ Cours < MM200 ({s['mm200']}€)",
             "DCFCE7" if s['above_mm200'] else "FFFBEB","DCFCE7" if s['above_mm200'] else "FFFBEB",True,"16A34A" if s['above_mm200'] else "D97706"),
            ("Zone DCF", "✅ En zone d'achat" if in_zone else f"— Hors zone (upside {s['upside']}%)",
             "DCFCE7" if in_zone else "F8FAFC","DCFCE7" if in_zone else "F8FAFC",True,"16A34A" if in_zone else "888888"),
        ]
        tri_score = sum(1 for row in tri if "✅" in row[1])
        two_col(doc, tri, f"Triptyque : {tri_score}/3 feux verts", "Lecture")
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

        # Thèse / Contra
        if s['thesis'] or s['contra']:
            t6 = doc.add_table(rows=2, cols=2); t6.style = 'Table Grid'
            for j,(h,hbg) in enumerate([("✅ Thèse haussière","16A34A"),("❌ Risques principaux","DC2626")]):
                c = t6.cell(0,j); setbg(c,hbg); rn(c.paragraphs[0],h,bold=True,color="FFFFFF",size=8.5)
            if s['thesis']:
                setbg(t6.cell(1,0),"F0FDF4")
                rn(t6.cell(1,0).paragraphs[0], s['thesis'][:280]+(("..." if len(s['thesis'])>280 else "")), size=8.5, color="1A4730")
            if s['contra']:
                setbg(t6.cell(1,1),"FFF5F5")
                rn(t6.cell(1,1).paragraphs[0], s['contra'][:280]+(("..." if len(s['contra'])>280 else "")), size=8.5, color="7C2C2C")
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

        # Track record
        if s['track']:
            tr_rows = []
            for y, p_txt, ok in s['track'][:3]:
                col_ = "16A34A" if ok=='ok' else "D97706" if ok=='partial' else "DC2626"
                bg_ = "F0FDF4" if ok=='ok' else "FFFBEB" if ok=='partial' else "FFF5F5"
                icon = "✓" if ok=='ok' else "~" if ok=='partial' else "✗"
                tr_rows.append((f"{icon} {y}", p_txt, bg_, bg_, ok=='ok', col_))
            two_col(doc, tr_rows, "Track record direction", "Résultat")
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

        # AVIS FINAL
        avis_verdict, avis_texte, avis_bg, avis_col = gen_avis(pos, s, pos['qarp'])
        avis_box(doc, avis_verdict, avis_texte, avis_bg, avis_col)

    # Footer
    doc.add_paragraph()
    pf = doc.add_paragraph(f"VAL.PEA · Revue de portefeuille · {date_str} · Cours de clôture · Non-conseil")
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.runs[0]; rf.italic = True; rf.font.size = Pt(8); rf.font.color.rgb = rgb("999999")

    doc.save(output_path)
    print(f"✅ Revue vendredi sauvegardée : {output_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 : SIGNAUX DIMANCHE — Grade A + zone + score ≥ 70
# ═══════════════════════════════════════════════════════════════════════════════
def generate_signaux_dimanche(html_content, output_path, date_str):
    stocks_all = parse_stocks(html_content)

    # Filtrer : Grade A + (in_zone ou dcf_zone) + score calculé ≥ 70
    signaux = []
    for s in stocks_all:
        if s['score'] != 'A':
            continue
        qarp = calc_qarp(s)
        if qarp['total'] < 70:
            continue
        if not (s['in_zone'] or s['dcf_zone']):
            continue
        s['qarp'] = qarp
        signaux.append(s)

    # Trier par score décroissant
    signaux.sort(key=lambda x: x['qarp']['total'], reverse=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    garde(doc, "SIGNAUX DU DIMANCHE",
          f"Dimanche {date_str}",
          f"{len(signaux)} opportunité(s) · Grade A · Zone · Score ≥ 70/100",
          "Protocole : vérifier les 5 étapes avant tout achat · Non-conseil")

    if not signaux:
        box(doc, "Aucun signal Grade A + zone + score ≥ 70 cette semaine. Patience — ne pas forcer.",
            "FFFBEB", "D97706", 11)
    else:
        para(doc, f"{len(signaux)} action(s) remplissent les 3 conditions : Grade A + en zone d'achat + score ≥ 70/100.",
             bold=True, color="16A34A", size=10, before=2, after=4)

        # Tableau récapitulatif
        headers = ["#","Ticker","Société","Score","Zone","R/R","Upside DCF","RSI","MM200","Pio."]
        t = doc.add_table(rows=len(signaux)+1, cols=10); t.style = 'Table Grid'
        for j,h in enumerate(headers):
            c = t.cell(0,j); setbg(c,"0F2540"); rn(c.paragraphs[0],h,bold=True,color="FFFFFF",size=8)
        for i,s in enumerate(signaux):
            for j,(text,bg_c,bold_,col_) in enumerate([
                (str(i+1),"F0FDF4",True,"16A34A"),
                (s['ticker'],"EFF6FF",True,"0F2540"),
                (s['name'][:14],"FAFAFA",False,"333333"),
                (f"{s['qarp']['total']}/100","DCFCE7",True,"16A34A"),
                (f"{s['el']}–{s['eh']}€","DCFCE7",False,"16A34A"),
                (f"{s['rr']}x","DCFCE7" if s['rr']>=1.5 else "FFFBEB",True,"16A34A" if s['rr']>=1.5 else "D97706"),
                (f"+{s['upside']}%","DCFCE7",False,"16A34A"),
                (f"{int(s['rsi'])}","DCFCE7" if s['rsi_ok'] else "FFFBEB",False,"16A34A" if s['rsi_ok'] else "D97706"),
                ("✅" if s['above_mm200'] else "⚠️","DCFCE7" if s['above_mm200'] else "FFFBEB",False,"16A34A" if s['above_mm200'] else "D97706"),
                (f"{int(s['pio'])}/9","DCFCE7" if s['pio']>=7 else "FAFAFA",False,"16A34A" if s['pio']>=7 else "333333"),
            ]):
                c = t.cell(i+1,j); setbg(c,bg_c); rn(c.paragraphs[0],text,bold=bold_,color=col_,size=8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Rappel protocole
        box(doc,
            "PROTOCOLE OBLIGATOIRE avant tout achat : "
            "(1) Grade A + Score ≥ 65  "
            "(2) Cours dans zone [el-eh] ou décote ≥ 12% sur DCF  "
            "(3) Triptyque 2/3 feux verts  "
            "(4) R/R ≥ 1.5x  "
            "(5) Taille max : Grade A = 8% du PEA",
            "EFF6FF", "1D4ED8", 9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # ── FICHES INDIVIDUELLES ──
        for idx, s in enumerate(signaux):
            if idx > 0:
                sep(doc)

            # Header
            th = doc.add_table(rows=1, cols=2); th.style = 'Table Grid'
            cl = th.cell(0,0); setbg(cl,"0F2540")
            p = cl.paragraphs[0]
            rn(p, "[A] ", bold=True, color="F0D080", size=12)
            rn(p, s['ticker'] + "  ", bold=True, color="F0D080", size=17)
            rn(p, s['name'][:24], color="AABBCC", size=9.5)
            p2 = cl.add_paragraph()
            rn(p2, f"{s['sector']}  ·  Score {s['qarp']['total']}/100  ·  "
                   f"Zone {s['el']}–{s['eh']}€  ·  R/R {s['rr']}x", color="AABBCC", size=8)
            cr = th.cell(0,1); setbg(cr,"065F46")
            p3 = cr.paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rn(p3, f"{s['qarp']['total']}/100  ", bold=True, color="F0D080", size=22)
            p4 = cr.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rn(p4, f"Cours : {s['price']}€  ·  Upside DCF : +{s['upside']}%", color="AABBCC", size=9)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

            # Zones
            in_zone = s['in_zone'] or s['dcf_zone']
            two_col(doc, [
                ("Cours actuel", f"{s['price']}€", "F8FAFC","F8FAFC",True,"0F2540"),
                ("Zone achat [el–eh]", f"{s['el']}€ – {s['eh']}€  ✅ EN ZONE", "DCFCE7","DCFCE7",True,"16A34A"),
                ("Stop Loss", f"{s['stop']}€  ({round((s['price']-s['stop'])/s['price']*100,1) if s['price'] else 0}% de marge)", "FEE2E2","FEE2E2",True,"DC2626"),
                ("Objectif 1", f"{s['o1']}€  (+{round((s['o1']-s['price'])/s['price']*100,1) if s['price'] else 0}%)", "DCFCE7","DCFCE7",True,"16A34A"),
                ("DCF Médian", f"{s['dcfm']}€  (upside +{s['upside']}%)", "EFF6FF","EFF6FF",True,"1D4ED8"),
                ("R/R", f"{s['rr']}x  {'✅ excellent' if s['rr']>=2 else '✅ favorable' if s['rr']>=1.5 else '⚠️ limite'}",
                 "DCFCE7","DCFCE7",True,"16A34A"),
            ], "Indicateur", "Valeur")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

            # Fondamentaux
            four_col(doc, [
                [("ROE","EFF6FF",True,"0F2540"),(f"{s['roe']}%","FAFAFA",False,"333333"),
                 ("Marge","EFF6FF",True,"0F2540"),(f"{s['margin']}%","FAFAFA",False,"333333")],
                [("FCF","EFF6FF",True,"0F2540"),(f"{s['fcf']}%","FAFAFA",False,"333333"),
                 ("Piotroski","EFF6FF",True,"0F2540"),(f"{int(s['pio'])}/9","DCFCE7",True,"16A34A")],
                [("Dette","EFF6FF",True,"0F2540"),(f"{s['debt']}x","FAFAFA",False,"333333"),
                 ("PE","EFF6FF",True,"0F2540"),(f"{s['pe']}x","FAFAFA",False,"333333")],
                [("RSI","EFF6FF",True,"0F2540"),(f"{int(s['rsi'])} {'✅' if s['rsi_ok'] else '⚠️'}","DCFCE7" if s['rsi_ok'] else "FFFBEB",False,"16A34A" if s['rsi_ok'] else "D97706"),
                 ("MM200","EFF6FF",True,"0F2540"),("✅" if s['above_mm200'] else "⚠️","DCFCE7" if s['above_mm200'] else "FFFBEB",False,"16A34A" if s['above_mm200'] else "D97706")],
            ], ["Ratio","Valeur","Ratio","Valeur"])
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

            # Score QARP détaillé
            q = s['qarp']
            two_col(doc, [
                ("Qualité (ROE + MOAT)", f"{q['q']}/20", "EFF6FF","FAFAFA",True,"0F2540"),
                ("Rentabilité (Marge + FCF)", f"{q['r']}/20", "EFF6FF","FAFAFA",True,"0F2540"),
                ("Bilan (Piotroski + Dette)", f"{q['b']}/20", "EFF6FF","FAFAFA",True,"0F2540"),
                ("Valorisation (Upside DCF)", f"{q['v']}/20", "EFF6FF","FAFAFA",True,"0F2540"),
                ("Momentum (Zone + MM200 + RSI)", f"{q['m']}/20", "EFF6FF","FAFAFA",True,"0F2540"),
                ("SCORE TOTAL", f"{q['total']}/100", "DCFCE7","DCFCE7",True,"16A34A"),
            ], "Dimension QARP", "Score")
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

            # Thèse / Contra
            if s['thesis'] or s['contra']:
                t6 = doc.add_table(rows=2, cols=2); t6.style = 'Table Grid'
                for j,(hh,hbg) in enumerate([("✅ Thèse haussière","16A34A"),("❌ Risques","DC2626")]):
                    c = t6.cell(0,j); setbg(c,hbg); rn(c.paragraphs[0],hh,bold=True,color="FFFFFF",size=8.5)
                if s['thesis']:
                    setbg(t6.cell(1,0),"F0FDF4")
                    rn(t6.cell(1,0).paragraphs[0], s['thesis'][:300]+(("..." if len(s['thesis'])>300 else "")), size=8.5, color="1A4730")
                if s['contra']:
                    setbg(t6.cell(1,1),"FFF5F5")
                    rn(t6.cell(1,1).paragraphs[0], s['contra'][:300]+(("..." if len(s['contra'])>300 else "")), size=8.5, color="7C2C2C")
                doc.add_paragraph().paragraph_format.space_after = Pt(3)

            # Avis signal
            avis_box(doc,
                     "✅ SIGNAL D'ACHAT VALIDÉ — Protocole à appliquer",
                     f"Grade A + Score {q['total']}/100 + en zone d'achat. "
                     f"R/R {s['rr']}x. Upside DCF +{s['upside']}%. "
                     f"Appliquer le protocole du dimanche avant tout passage d'ordre. "
                     f"Taille max : 8% du PEA. Stop loss : {s['stop']}€.",
                     "DCFCE7", "16A34A")

    # Footer
    doc.add_paragraph()
    pf = doc.add_paragraph(f"VAL.PEA · Signaux du dimanche · {date_str} · Non-conseil en investissement")
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.runs[0]; rf.italic = True; rf.font.size = Pt(8); rf.font.color.rgb = rgb("999999")

    doc.save(output_path)
    print(f"✅ Signaux dimanche sauvegardés : {output_path}")


# ─── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    html_file = 'index.html'
    if not os.path.exists(html_file):
        print(f"ERROR: {html_file} not found"); sys.exit(1)

    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()

    now = datetime.now()
    date_str = now.strftime('%d %B %Y').replace(
        'January','janvier').replace('February','février').replace('March','mars').replace(
        'April','avril').replace('May','mai').replace('June','juin').replace(
        'July','juillet').replace('August','août').replace('September','septembre').replace(
        'October','octobre').replace('November','novembre').replace('December','décembre')

    mode = sys.argv[1] if len(sys.argv) > 1 else 'both'

    if mode in ('vendredi', 'friday', 'both'):
        fname = f"VAL_PEA_Revue_{now.strftime('%Y%m%d')}.docx"
        generate_revue_vendredi(html_content, fname, date_str)

    if mode in ('dimanche', 'sunday', 'both'):
        fname = f"VAL_PEA_Signaux_{now.strftime('%Y%m%d')}.docx"
        generate_signaux_dimanche(html_content, fname, date_str)

    print("DONE")
